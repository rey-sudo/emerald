from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt
from pydantic import BaseModel, Field, TypeAdapter


# =========================
# SCHEMA INMUTABLE
# =========================

class QuestionItem(BaseModel):
    question: str = Field(
        ...,
        description="Enunciado de la pregunta tipo caso con mención de artículo y normativa"
    )

    options: List[str] = Field(
        ...,
        min_length=4,
        max_length=4,
        description="Lista de opciones posibles"
    )

    correct: int = Field(
        ...,
        ge=0,
        le=3,
        description="Índice de la opción correcta"
    )

    explanation: str = Field(
        ...,
        description="Explicación completa de la respuesta"
    )


questions_adapter = TypeAdapter(List[QuestionItem])


# =========================
# CONFIG
# =========================

OPTION_LABELS = ("A", "B", "C", "D")


# =========================
# LOAD + VALIDATION
# =========================

def load_questions(json_path: str | Path) -> List[QuestionItem]:
    """
    Carga y valida preguntas desde JSON.
    """
    json_path = Path(json_path)

    raw_json = json_path.read_text(encoding="utf-8")

    return questions_adapter.validate_json(raw_json)


# =========================
# WORD RENDERING
# =========================

def build_document(questions: List[QuestionItem]) -> Document:
    """
    Construye el documento Word.
    """
    doc = Document()

    title = doc.add_heading("Banco de Preguntas", level=1)
    title.runs[0].font.size = Pt(20)

    for index, item in enumerate(questions, start=1):
        add_question_block(doc, index, item)

    return doc


def add_question_block(
    doc: Document,
    number: int,
    item: QuestionItem
) -> None:
    """
    Renderiza una pregunta completa.
    """

    # Pregunta
    p_question = doc.add_paragraph()

    run = p_question.add_run(f"{number}. {item.question}")
    run.bold = True
    run.font.size = Pt(12)

    # Opciones
    for label, option in zip(OPTION_LABELS, item.options):
        doc.add_paragraph(
            f"{label}. {option}",
            style="List Bullet"
        )

    # Separador visual
    doc.add_paragraph("—" * 35)


# =========================
# EXPORT
# =========================

def export_questions_to_word(
    input_json: str | Path,
    output_docx: str | Path = "preguntas.docx"
) -> Path:
    """
    Pipeline completo:
    JSON -> Validación -> Word
    """

    questions = load_questions(input_json)

    doc = build_document(questions)

    output_path = Path(output_docx)

    doc.save(output_path)

    return output_path


# =========================
# MAIN
# =========================

if __name__ == "__main__":

    output_file = export_questions_to_word(
        input_json="preguntas.json",
        output_docx="preguntas.docx"
    )

    print(f"Documento generado: {output_file}")